#!/usr/bin/env python3
"""
Full Contract Test Script for W2 Clause Review Workflow
Tests a complete PSA document by extracting clauses and sending each to the W2 endpoint.
"""

import json
import re
import requests
import time
import uuid
from datetime import datetime
from docx import Document
from pathlib import Path

# Configuration
W2_ENDPOINT = "https://mmenendeza.app.n8n.cloud/webhook/clause-review-rag"
OUTPUT_DIR = Path("./config/test_results")
TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M%S")

def clean_text(text: str) -> str:
    """Clean text by normalizing Unicode characters common in Word documents."""
    replacements = {
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote  
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u00a0': ' ',  # Non-breaking space
        '\t': ' ',      # Tab
        '\r': '',       # Carriage return
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Normalize whitespace
    text = ' '.join(text.split())
    return text

def extract_clauses_from_docx(filepath: str) -> list:
    """Extract clause-like sections from a DOCX file.
    Groups heading with following paragraphs until next heading.
    """
    doc = Document(filepath)
    clauses = []
    
    # Patterns that indicate a section heading (standalone or embedded)
    heading_patterns = [
        r'^(EFFECTIVE DATE|PARTIES|PRODCO|SERVICES|RIGHTS|FEES|ENTITLEMENTS|CREDIT|'
        r'REPRESENTATIONS|WARRANTIES|INDEMNITY|MISCELLANEOUS|CONDITIONS|DAMAGES|'
        r'INJUNCTIVE|ASSIGNMENT|DATA PROTECTION|TAX|GOVERNING LAW|JURISDICTION|'
        r'CONFIDENTIAL|TERMINATION|LIABILITY|PAYMENT|INSURANCE|AUDIT|DISPUTE|'
        r'FORCE MAJEURE|SURVIVAL|REMEDY|REVERSION)[:\;\s]'
    ]
    
    current_heading = None
    current_text_parts = []
    
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text or len(text) < 10:  # Skip very short
            continue
        
        # Check if this is a new heading
        is_heading = False
        new_heading = None
        for pattern in heading_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                is_heading = True
                new_heading = match.group(1).upper()
                break
        
        if is_heading:
            # Save previous section if we have content
            if current_heading and current_text_parts:
                full_text = ' '.join(current_text_parts)
                if len(full_text) > 50:  # Only save substantial sections
                    clauses.append({
                        "heading": current_heading,
                        "text": full_text,
                        "section": current_heading
                    })
            
            # Start new section
            current_heading = new_heading
            current_text_parts = [text]  # Include heading line in text
        else:
            # Append to current section
            if current_heading:
                current_text_parts.append(text)
            else:
                # First paragraph before any heading - might be preamble
                if len(text) > 100:
                    clauses.append({
                        "heading": text[:40] + "...",
                        "text": text,
                        "section": "PREAMBLE"
                    })
    
    # Don't forget the last section
    if current_heading and current_text_parts:
        full_text = ' '.join(current_text_parts)
        if len(full_text) > 50:
            clauses.append({
                "heading": current_heading,
                "text": full_text,
                "section": current_heading
            })
    
    # Also extract from tables (but with more context)
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            row_text = ' | '.join([c for c in cells if c])
            if row_text:
                table_text.append(row_text)
        
        full_text = '\n'.join(table_text)
        if len(full_text) > 100:
            # Try to get a heading from first row
            first_row = table_text[0] if table_text else "Table"
            heading = first_row[:50] if len(first_row) > 50 else first_row
            clauses.append({
                "heading": heading,
                "text": full_text,
                "section": "TABLE"
            })
    
    return clauses

def send_clause_to_w2(clause: dict, idx: int, run_id: str, doc_id: str) -> dict:
    """Send a single clause to W2 endpoint and return response."""
    payload = {
        "clause_instance_id": f"psa-clause-{idx:03d}",
        "clause_id": f"clause-{idx}",
        "clause_text": clause["text"][:4000],  # Truncate very long clauses
        "heading": clause.get("heading", f"Clause {idx}"),
        "run_id": run_id,
        "document_id": doc_id,
        "playbook_id": "PB_DSA_V1"
    }
    
    try:
        response = requests.post(
            W2_ENDPOINT,
            headers={"Content-Type": "application/json"},
            json=payload,
            timeout=120
        )
        if response.status_code == 200:
            try:
                return response.json()
            except:
                return {"error": "Invalid JSON response", "raw": response.text[:200]}
        else:
            return {"error": f"HTTP {response.status_code}", "raw": response.text[:200]}
    except requests.exceptions.Timeout:
        return {"error": "Timeout"}
    except Exception as e:
        return {"error": str(e)}

def run_full_contract_test(docx_path: str):
    """Run full contract test and generate report."""
    print(f"\n{'='*60}")
    print(f"FULL CONTRACT TEST - W2 Clause Review")
    print(f"Document: {docx_path}")
    print(f"Endpoint: {W2_ENDPOINT}")
    print(f"Timestamp: {TIMESTAMP}")
    print(f"{'='*60}\n")
    
    # Extract clauses
    print("Extracting clauses from document...")
    clauses = extract_clauses_from_docx(docx_path)
    print(f"Found {len(clauses)} clauses\n")
    
    if not clauses:
        print("ERROR: No clauses extracted from document")
        return
    
    # Prepare test run
    run_id = str(uuid.uuid4())
    doc_id = str(uuid.uuid4())
    results = []
    
    # Warm-up ping
    print("Warming up endpoint...")
    requests.post(W2_ENDPOINT, 
                  json={"clause_instance_id": "warmup", "clause_text": "test", 
                        "run_id": run_id, "document_id": doc_id},
                  timeout=60)
    time.sleep(2)
    
    # Process each clause
    family_counts = {}
    decision_counts = {}
    
    for idx, clause in enumerate(clauses, 1):
        heading = clause.get("heading", "")[:40]
        print(f"[{idx:3d}/{len(clauses)}] {heading}...", end=" ", flush=True)
        
        result = send_clause_to_w2(clause, idx, run_id, doc_id)
        
        family = result.get("detected_family", "ERROR")
        decision = result.get("decision", "ERROR")
        has_spec = "PB:" in result.get("_internal", {}).get("rule_id", "")
        
        print(f"{family} | {decision} | {'✓' if has_spec else '○'}")
        
        # Track stats
        family_counts[family] = family_counts.get(family, 0) + 1
        decision_counts[decision] = decision_counts.get(decision, 0) + 1
        
        results.append({
            "idx": idx,
            "heading": clause.get("heading", ""),
            "text_preview": clause["text"][:100] + "...",
            "family": family,
            "decision": decision,
            "has_playbook_spec": has_spec,
            "full_response": result
        })
        
        # Rate limiting
        time.sleep(1.5)
    
    # Generate summary
    print(f"\n{'='*60}")
    print("TEST SUMMARY")
    print(f"{'='*60}")
    print(f"\nTotal clauses: {len(clauses)}")
    print("\nFamily Distribution:")
    for family, count in sorted(family_counts.items(), key=lambda x: -x[1]):
        pct = count/len(clauses)*100
        print(f"  {family}: {count} ({pct:.1f}%)")
    
    print("\nDecision Distribution:")
    for decision, count in sorted(decision_counts.items(), key=lambda x: -x[1]):
        pct = count/len(clauses)*100
        print(f"  {decision}: {count} ({pct:.1f}%)")
    
    # Save results
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_file = OUTPUT_DIR / f"full_contract_test_{TIMESTAMP}.json"
    
    report = {
        "timestamp": TIMESTAMP,
        "document": docx_path,
        "total_clauses": len(clauses),
        "family_distribution": family_counts,
        "decision_distribution": decision_counts,
        "results": results
    }
    
    with open(output_file, "w") as f:
        json.dump(report, f, indent=2)
    
    print(f"\nResults saved to: {output_file}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        # Default test document
        docx_path = "contratos de test/PSA Principal Terms Fallback Guide (20 October 2025) - Prueba 1.docx"
    
    run_full_contract_test(docx_path)
